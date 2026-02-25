from fastapi import FastAPI, HTTPException, Request
from pydantic import BaseModel, Field, validator
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from groq import Groq
from dotenv import load_dotenv
from transformers import (
    AutoConfig,
    AutoTokenizer,
    AutoModelForSeq2SeqLM,
    AutoModelForCausalLM,
)
from slowapi import Limiter
from slowapi.util import get_remote_address
from pathlib import Path
import uvicorn
import torch
import os
import re
import time
from pathlib import Path

# ===============================
# APP SETUP
# ===============================
load_dotenv()

BASE_DIR = Path(__file__).resolve().parent.parent
UPLOAD_DIR = (BASE_DIR / "uploads").resolve()

app = FastAPI()
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter

# ===============================
# CONFIG
# ===============================
HF_GENERATION_MODEL = os.getenv("HF_GENERATION_MODEL", "google/flan-t5-small")
LLM_GENERATION_TIMEOUT = int(os.getenv("LLM_GENERATION_TIMEOUT", "30"))
SESSION_TIMEOUT = 3600

sessions = {}

embedding_model = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)

SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".txt", ".md"]

# -------------------------------------------------------------------
# DOCUMENT LOADERS
# -------------------------------------------------------------------
def load_pdf(file_path: str) -> list[Document]:
    loader = PyPDFLoader(file_path)
    return loader.load()

def load_docx(file_path: str) -> list[Document]:
    import docx
    doc = docx.Document(file_path)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend([p.text.strip() for p in cell.paragraphs if p.text.strip()])
    full_text = "\n".join(texts)
    return [Document(page_content=full_text, metadata={"source": file_path})]

def load_txt(file_path: str) -> list[Document]:
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    return [Document(page_content=content, metadata={"source": file_path})]

def load_document(file_path: str) -> list[Document]:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext in (".txt", ".md"):
        return load_txt(file_path)
    else:
        raise ValueError(f"Unsupported format: {ext}")

splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)

# -------------------------------------------------------------------
# TEXT NORMALIZATION
# -------------------------------------------------------------------
def normalize_spaced_text(text: str) -> str:
    pattern = r"\b(?:[A-Za-z] ){2,}[A-Za-z]\b"
    return re.sub(pattern, lambda m: m.group(0).replace(" ", ""), text)


def normalize_answer(text: str) -> str:
    """
    Post-processes the LLM-generated answer.
    """
    text = normalize_spaced_text(text)
    text = re.sub(r"^(Answer[^:]*:|Context:|Question:)\s*", "", text, flags=re.I)
    return text.strip()


# ===============================
# DOCUMENT LOADERS
# ===============================
def load_pdf(file_path: str):
    return PyPDFLoader(file_path).load()


def load_txt(file_path: str):
    with open(file_path, "r", encoding="utf-8") as f:
        return [Document(page_content=f.read())]


def load_docx(file_path: str):
    doc = docx.Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(page_content=text)]


def load_document(file_path: str):
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext in [".txt", ".md"]:
        return load_txt(file_path)
    else:
        raise ValueError("Unsupported file format")


# ===============================
# MODEL LOADING
# ===============================
def load_generation_model():
    global generation_model, generation_tokenizer, generation_is_encoder_decoder

    if generation_model:
        return generation_tokenizer, generation_model, generation_is_encoder_decoder

    config = AutoConfig.from_pretrained(HF_GENERATION_MODEL)
    generation_is_encoder_decoder = bool(config.is_encoder_decoder)

    generation_tokenizer = AutoTokenizer.from_pretrained(HF_GENERATION_MODEL)

    if generation_is_encoder_decoder:
        generation_model = AutoModelForSeq2SeqLM.from_pretrained(HF_GENERATION_MODEL)
    else:
        generation_model = AutoModelForCausalLM.from_pretrained(HF_GENERATION_MODEL)

    if torch.cuda.is_available():
        generation_model = generation_model.to("cuda")

    generation_model.eval()
    return generation_tokenizer, generation_model, generation_is_encoder_decoder


def generate_response(prompt: str, max_new_tokens: int):
    tokenizer, model, is_enc = load_generation_model()
    device = next(model.parameters()).device

    encoded = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=2048)
    encoded = {k: v.to(device) for k, v in encoded.items()}

    output = model.generate(
        **encoded,
        max_new_tokens=max_new_tokens,
        do_sample=False,
        pad_token_id=tokenizer.pad_token_id or tokenizer.eos_token_id,
    )

    if is_enc:
        return tokenizer.decode(output[0], skip_special_tokens=True)

    return tokenizer.decode(
        output[0][encoded["input_ids"].shape[1]:],
        skip_special_tokens=True,
    )


# ===============================
# REQUEST MODELS
# ===============================
class DocumentPath(BaseModel):
    filePath: str
    session_id: str


class AskRequest(BaseModel):
    question: str = Field(..., min_length=1)
    session_id: str
    history: list = []

    @validator("question")
    def validate_question(cls, v):
        if not v.strip():
            raise ValueError("Empty question")
        return v.strip()


class SummarizeRequest(BaseModel):
    session_id: str
    pdf: str | None = None



class CompareRequest(BaseModel):
    session_id: str

# -------------------------------------------------------------------
# SESSION CLEANUP
# -------------------------------------------------------------------


def cleanup_expired_sessions():
    now = time.time()
    expired = [k for k, v in sessions.items()
               if now - v["last"] > SESSION_TIMEOUT]
    for k in expired:
        del sessions[k]


# ===============================
# PROCESS DOCUMENT
# ===============================
@app.post("/process")
@limiter.limit("15/15 minutes")
def process_doc(request: Request, data: DocumentPath):
    cleanup_sessions()

    if not os.path.exists(data.filePath):
        raise HTTPException(404, "File not found")

    docs = load_document(data.filePath)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
    )
    chunks = splitter.split_documents(docs)

    vectorstore = FAISS.from_documents(chunks, embedding_model)

    if data.session_id in sessions:
        # Append to existing vectorstore
        sessions[data.session_id]["vectorstore"].add_documents(chunks)
        sessions[data.session_id]["last_accessed"] = time.time()
    else:
        # Create new vectorstore
        sessions[data.session_id] = {
            "vectorstore": FAISS.from_documents(chunks, embedding_model),
            "last_accessed": time.time(),
        }

    return {"message": "Processed successfully"}


# ===============================
# ASK
# ===============================
@app.post("/ask")
@limiter.limit("60/15 minutes")
def ask(request: Request, data: AskRequest):
    cleanup_sessions()

    session = sessions.get(data.session_id)
    if not session:
        return {"answer": "Session expired", "confidence_score": 0}

    vectorstore = session["vectorstore"]

    docs = vectorstore.similarity_search_with_score(data.question, k=4)

    docs_with_scores = vectorstore.similarity_search_with_score(question, k=4)
    if not docs_with_scores:
        return {"answer": "No relevant context found in the uploaded document.", "confidence_score": 0}

    # Convert FAISS scores to cosine similarities
    scored = [(doc, score, faiss_score_to_cosine_sim(score)) for doc, score in docs_with_scores]
    similarities = [sim for _, _, sim in scored]
    confidence = compute_confidence([score for _, score, _ in scored])

    # Relevance threshold check
    top3_avg_sim = sum(sorted(similarities, reverse=True)[:3]) / 3 if similarities else 0
    if top3_avg_sim < RELEVANCE_THRESHOLD:
        return {
            "answer": "I cannot answer this question based on the uploaded document. It appears unrelated.",
            "confidence_score": confidence
        }

    relevant_docs = [doc for doc, _, sim in scored if sim >= RELEVANCE_THRESHOLD]
    context = "\n\n".join([d.page_content for d in relevant_docs])

    Question:
    {data.question}

    user_prompt = CCC_PROMPT.format(
        context=context,
        question=question_with_history,
    )

    answer = generate_response(prompt, 150)

    session["last"] = time.time()

    return {"answer": normalize_answer(answer), "confidence_score": 85}

@app.post("/summarize")
@limiter.limit("15/15 minutes")
async def summarize_pdf(request: Request, data: SummarizeRequest):
    cleanup_expired_sessions()
    session_id = data.session_id
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found or expired")

    sessions[session_id]["last_accessed"] = time.time()
    vectorstore = sessions[session_id]["vectorstore"]

    # Extract all text from the vectorstore for summarization
    all_docs = list(vectorstore.docstore._dict.values())
    full_text = "\n".join([doc.page_content for doc in all_docs])

    system_prompt = (
        "You are a document summarization assistant.\n"
        "Rules:\n"
        "1. Summarize in 6-8 concise bullet points.\n"
        "2. Clearly state: who received the certificate/document, what it is for, "
        "which organization issued it, who authorized it, and the date.\n"
        "3. Use proper Title Case for names. Return clean, readable text.\n"
        "4. Use ONLY the information in the provided context."
    )

    user_prompt = f"Content:\n{full_text[:4000]}\n\nSummary (bullet points):"

    return {"summary": normalize_answer(summary)}


@app.post("/compare")
@limiter.limit("15/15 minutes")
async def compare_documents(request: Request, data: CompareRequest):
    """
    Compare multiple documents within a session.
    Groups chunks by their source metadata and generates a comparative summary.
    """
    cleanup_expired_sessions()
    session_id = data.session_id
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found or expired")

    sessions[session_id]["last_accessed"] = time.time()
    vectorstore = sessions[session_id]["vectorstore"]

    # Extract all documents from the FAISS internal docstore
    try:
        all_docs = list(vectorstore.docstore._dict.values())
    except Exception:
        all_docs = []

    if not all_docs:
        return {"comparison": "No documents found to compare."}

    # Group chunks by their source file
    docs_by_source = {}
    for doc in all_docs:
        source = doc.metadata.get("source", "Unknown")
        if source not in docs_by_source:
            docs_by_source[source] = []
        docs_by_source[source].append(doc.page_content)

    sources = list(docs_by_source.keys())
    if len(sources) < 2:
        return {"comparison": "Please upload at least two documents to generate a comparison."}

    # Prepare context for comparison
    comparison_context = ""
    for i, source in enumerate(sources):
        filename = Path(source).name
        content_sample = "\n".join(docs_by_source[source])[:1500]
        comparison_context += f"--- Document {i+1}: {filename} ---\n{content_sample}\n\n"

    system_prompt = "You are a professional analyst comparing multiple documents."
    user_prompt = f"""
    Please compare the following documents:

    {comparison_context}

    Instructions:
    - Identify key similarities and differences.
    - Format as a structured comparison (bullet points/sections).
    - Be objective and concise.
    """

    comparison_result = generate_response(
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        max_tokens=800
    )
    return {"comparison": comparison_result}


# -------------------------------------------------------------------
# START SERVER
# -------------------------------------------------------------------
if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=5000, reload=True)